from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("13节点全量学习计划更新任务书.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5E6E82"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
LIGHT_GRAY = "F2F4F7"


def set_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
                element = margins.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    margins.append(element)
                element.set(qn("w:w"), value)
                element.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9.5, bold=True, color=INK)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for index, line in enumerate(str(value).split("\n")):
                if index:
                    p.add_run().add_break()
                r = p.add_run(line)
                set_font(r, size=9.2, color=INK)
    set_table_geometry(table, widths)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}[level], bold=True,
             color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead:
        r = p.add_run(bold_lead)
        set_font(r, bold=True, color=INK)
    r = p.add_run(text)
    set_font(r, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, color=INK)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    set_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    return table


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("学习计划能力节点闭环 | 更新任务书")
    set_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("内部实施文件 | 2026-08-24")
    set_font(r, size=9, color=MUTED)


def build():
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("实施任务书")
    set_font(r, size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("13 节点全量学习计划与覆盖推进机制更新")
    set_font(r, size=23, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("适用范围：RAG 工程链路培训；设计应可复用于其他知识图谱方向")
    set_font(r, size=12, color=MUTED)

    add_table(doc, ["项目项", "约定"], [
        ("文档版本", "v1.0"),
        ("发布日期", "2026-08-24"),
        ("目标", "让 13 个能力节点在多轮学习中可追踪、可解释、不会长期遗漏。"),
        ("执行边界", "不自动发起下一轮生成；由系统推荐，学习者确认后创建任务。"),
        ("成功定义", "每个节点都具有持久化推进状态、选点原因和可验证的覆盖轨迹。"),
    ], [2700, 6660])

    add_heading(doc, "1. 背景与问题", 1)
    add_body(doc, "当前系统已能基于掌握度、资源发布记录和前置关系推荐节点，但“节点是否已安排、是否完成学习、是否等待验证”主要由多个数据源推导。长期运行时，低优先级且不属于薄弱项的节点可能不断被补强任务挤出。")
    add_callout(doc, "本次决策", "新增独立的课程推进记录和覆盖欠债机制；它们不替代现有客观掌握度，而是补齐学习过程与多轮调度能力。")
    add_heading(doc, "1.1 本次要解决的问题", 2)
    for text in [
        "13 个节点没有统一的可审计推进状态，难以解释某节点为何尚未进入学习。",
        "资源发布可代表“已接触”，但不能严格表达“已计划、已学习、待验证、已完成或需巩固”。",
        "仅依赖当前优先级时，满足前置但优先级较低的节点缺少等待时间保障。",
        "生成任务失败、重试或重复提交时，节点状态需要具备可恢复与幂等语义。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "1.2 非目标", 2)
    for text in [
        "不把资源发布等同于掌握；是否掌握仍以服务端正式测评和证据为准。",
        "不绕过知识图谱前置关系，也不因欠债机制强制越级学习。",
        "不在未获学习者确认的情况下自动创建、自动执行下一轮生成任务。",
        "不改变现有生成、反馈、资源发布 API 的既有字段语义；新增字段须向前兼容。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2. 目标架构与状态模型", 1)
    add_body(doc, "保留两条独立但联动的数据线：客观掌握度回答“学得如何”，课程推进状态回答“学习流程走到哪里”。调度器同时读取两者。")
    add_table(doc, ["数据线", "来源", "用途", "权威性"], [
        ("掌握度", "诊断、正式测评、学习尝试证据", "判断弱项、退步、已掌握", "服务端客观证据"),
        ("课程推进", "选点、任务、发布、反馈事件", "保证覆盖、恢复中断、解释推荐", "持久化课程记录"),
        ("资源覆盖", "已发布资源的目标节点", "证明节点至少获得过学习材料", "发布成功后生效"),
    ], [1600, 2500, 3100, 2160])
    add_heading(doc, "2.1 课程推进状态", 2)
    add_table(doc, ["状态", "含义", "进入条件", "主要退出条件"], [
        ("unplanned", "未计划", "画像初始化或任务回退", "被本轮原子锁定"),
        ("scheduled", "已计划", "生成任务已创建", "资源发布；或任务终态失败/取消"),
        ("exposed", "已学习", "至少一个目标资源已发布", "进入待验证；或补强流程"),
        ("verification_pending", "待验证", "已学习但尚无本轮正式测评结论", "测评达标或不足"),
        ("completed", "已完成", "资源发布且正式测评达标", "后续客观证据显示需要巩固"),
        ("reinforcement_due", "需巩固", "测评薄弱、退步或策略判定需要补强", "补强资源发布后重新待验证"),
    ], [1500, 1800, 3000, 3060])
    add_callout(doc, "状态约束", "状态流转必须由服务端用例执行；前端只能读取、确认建议或提交正式反馈，不得直接改写课程推进状态。")

    add_heading(doc, "3. 数据库与迁移任务", 1)
    add_body(doc, "新增表名建议为 learner_curriculum_nodes。采用可回滚、向前兼容迁移，不重命名或删除现有掌握度、资源、生成任务表。")
    add_table(doc, ["字段", "类型/约束", "说明"], [
        ("learner_id", "TEXT，联合主键", "学习者标识"),
        ("knowledge_base_id", "TEXT，联合主键", "知识图谱/学习方向标识"),
        ("skill_node_id", "TEXT，联合主键", "能力节点 ID；不可使用展示名称"),
        ("progress_status", "TEXT，非空", "六态课程推进状态"),
        ("wait_rounds", "INTEGER，非负，默认 0", "前置满足后未进入本轮的连续轮数"),
        ("scheduled_run_id", "TEXT，可空", "当前或最近锁定该节点的生成 Run"),
        ("published_resource_count", "INTEGER，非负，默认 0", "只统计已发布的目标资源"),
        ("verified_attempt_count", "INTEGER，非负，默认 0", "正式测评次数"),
        ("last_scheduled_at / last_published_at / last_verified_at", "UTC 时间，可空", "审计与排序依据"),
        ("row_version", "INTEGER，非负", "CAS/幂等更新版本"),
        ("created_at / updated_at", "UTC 时间", "审计字段"),
    ], [2860, 2500, 4000])
    add_heading(doc, "3.1 迁移验收", 2)
    for text in [
        "对已有学习者按其知识库图谱补齐全部节点记录；RAG 样例必须为 13 条。",
        "迁移可重复运行，不产生重复记录；新旧 SQLite 数据库均可启动。",
        "已有已发布资源可回填 published_resource_count，但不得把未发布草稿计为学习覆盖。",
        "节点已被正式测评且达标时，可回填 completed；证据不足时保守回填 exposed 或 verification_pending。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4. 多轮选点与覆盖欠债策略", 1)
    add_body(doc, "每轮最多选 3 个节点。选点器仅从前置条件满足的节点中选取，按补强需求、覆盖欠债、图谱影响范围和学习者显式选择综合排序。")
    add_heading(doc, "4.1 候选池与排序", 2)
    add_table(doc, ["候选池", "准入条件", "排序与名额规则"], [
        ("补强池", "reinforcement_due，且前置已满足", "客观弱项、退步优先；最多占 2 个默认名额"),
        ("新知识池", "unplanned，且前置已覆盖/完成", "覆盖欠债、下游影响、稳定顺序；默认至少保留 1 个名额"),
        ("待验证池", "exposed 或 verification_pending", "不直接替代生成名额；用于提示优先安排测评"),
        ("阻塞池", "存在未完成前置", "不可选；返回被阻塞节点与具体前置原因"),
    ], [1800, 3100, 4460])
    add_heading(doc, "4.2 覆盖欠债规则", 2)
    for text in [
        "仅在节点前置关系已满足且节点仍为 unplanned 时累计 wait_rounds；被前置阻塞期间不累计。",
        "每完成一轮“用户确认且成功创建”的选点结算：未选中的合格节点 wait_rounds +1；被选中节点归零。",
        "默认阈值为 3 轮。达到阈值的节点进入强制候选池，下一轮至少分配 1 个新知识名额；若无新知识名额或存在硬性前置阻塞，则记录原因但不越级。",
        "欠债仅为排序加分并设上限，避免无限放大；建议上限为 5。所有排序须具有稳定的 node_id 兜底顺序。",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "默认配额", "建议 0-2 个需巩固节点 + 1 个未覆盖节点；若没有需巩固节点，则从合格的新知识池补足至最多 3 个。")

    add_heading(doc, "5. 生成、发布与反馈闭环", 1)
    add_number(doc, "推荐：服务端计算候选节点、状态、等待轮数和推荐原因；前端展示推荐与阻塞原因。")
    add_number(doc, "确认：学习者确认 1-3 个节点与资源类型；服务端校验快照、前置、上限和状态版本。")
    add_number(doc, "锁定：在创建 Generation Job 的同一事务中，将节点更新为 scheduled，并把 target_skill_nodes、计划快照和排序原因写入请求 payload。")
    add_number(doc, "生成：规划、检索、资源规格、内容生成和审核均读取 frozen target_skill_nodes。资源规格必须保留节点 ID。")
    add_number(doc, "发布：资源通过发布门后，按 resource.knowledge_points 中的目标节点更新 published_resource_count，并变更为 exposed / verification_pending。")
    add_number(doc, "反馈：正式测评更新客观掌握度与课程推进状态；系统重新计算下一轮建议，但不自动建任务。")
    add_heading(doc, "5.1 失败与重试语义", 2)
    add_table(doc, ["事件", "必须行为", "禁止行为"], [
        ("任务创建幂等重放", "复用原 run_id 与节点锁定；不得重复增加 wait_rounds", "创建第二份计划记录或重复扣减名额"),
        ("生成失败/取消", "仅在该节点未有已发布资源时回退至 unplanned；写入失败原因", "将失败草稿计为已学习"),
        ("部分资源发布", "只更新实际发布且目标节点明确的记录", "整批任务一律标记为已完成"),
        ("过期快照确认", "拒绝并要求前端刷新推荐", "按旧状态继续创建任务"),
    ], [2000, 4050, 3310])

    add_heading(doc, "6. 接口与前端更新", 1)
    add_table(doc, ["位置", "更新要求", "兼容性"], [
        ("能力画像接口", "增加 curriculum_progress：总数、各状态计数、节点状态、wait_rounds、选点 rank 与原因。", "保留既有 mastery、priority 和节点字段。"),
        ("下一轮推荐接口", "返回 recommended_node_ids、强制候选标记、阻塞原因、计划快照版本。", "保留 reinforce_weakness 与 learn_new_knowledge。"),
        ("生成任务创建", "写入 curriculum_snapshot 与 target_skill_nodes；返回已锁定节点。", "继续支持显式 target_skill_nodes。"),
        ("反馈确认接口", "确认前校验快照与课程记录版本；成功后触发下一轮候选重算。", "仍由用户确认，不自动生成。"),
        ("反馈页面", "展示“已完成 x/13、待学习、待验证、需巩固”；标注推荐/欠债/前置阻塞原因。", "保持每轮最多选择 3 个节点。"),
    ], [1900, 4650, 2810])

    add_heading(doc, "7. 任务分解与验收标准", 1)
    add_table(doc, ["编号", "任务", "主要产出", "验收标准", "依赖"], [
        ("T1", "领域模型与迁移", "课程推进 DTO、SQLite 迁移、仓储接口", "13 节点初始化；迁移可重跑；无破坏性变更", "无"),
        ("T2", "状态机服务", "原子状态流转、CAS/幂等策略、审计原因", "合法流转通过；非法流转拒绝；失败可恢复", "T1"),
        ("T3", "选点与欠债调度", "候选池、稳定排序、配额与阈值规则", "前置不越级；3 轮欠债节点进入推荐；每轮 <=3", "T1,T2"),
        ("T4", "生成与发布联动", "任务锁定、target 快照、发布回写", "目标节点贯穿任务/资源；未发布不计覆盖", "T2,T3"),
        ("T5", "反馈与测评联动", "测评后的推进状态更新、刷新推荐", "达标/薄弱正确转态；不自动创建任务", "T2,T4"),
        ("T6", "API 与前端展示", "进度概览、推荐原因、阻塞说明", "可看清 13 节点进度；确认校验可用", "T3,T5"),
        ("T7", "回归与文档", "测试矩阵、API/架构文档、迁移说明", "全部约定测试通过；变更可追溯", "T1-T6"),
    ], [550, 1650, 2700, 3000, 1460])
    add_heading(doc, "7.1 建议执行顺序", 2)
    for text in [
        "先完成 T1/T2 并提供独立单元测试；状态机未稳定前不得改造前端。",
        "再完成 T3，使用固定 13 节点图谱 fixture 验证多轮调度。",
        "随后接入 T4/T5，重点覆盖发布、失败、重试和正式测评事件。",
        "最后完成 T6/T7，并在迁移后 SQLite 数据上执行端到端回归。",
    ]:
        add_number(doc, text)

    add_heading(doc, "8. 测试矩阵", 1)
    add_table(doc, ["层级", "必须覆盖的场景"], [
        ("单元", "六态流转、非法流转、CAS 冲突、wait_rounds 上限、稳定排序、最多 3 节点。"),
        ("服务集成", "任务创建锁定、幂等重放、失败回退、发布回写、正式测评转态、过期快照拒绝。"),
        ("迁移", "新库初始化、已有学习者回填、重复迁移、已发布与未发布资源的区分。"),
        ("端到端", "13 节点图谱多轮运行，前置顺序正确，所有节点最终至少曝光一次。"),
        ("前端", "进度计数、推荐原因、阻塞说明、欠债标记、选择上限和确认后的状态刷新。"),
        ("回归", "现有问卷、诊断、生成、发布、反馈、资源阅读与 API 状态码不退化。"),
    ], [1800, 7560])
    add_callout(doc, "关键验收演示", "以 13 节点 RAG 图谱创建新画像；持续执行“推荐 -> 用户确认 -> 发布 -> 测评”循环。演示中应能看到每个节点的推进状态、等待轮数与原因，且最终所有节点均至少产生一次已发布资源。")

    add_heading(doc, "9. 风险与实施约束", 1)
    add_table(doc, ["风险", "控制措施"], [
        ("把发布误当掌握", "推进状态和客观掌握度分离；completed 必须依赖正式测评。"),
        ("欠债造成越级", "只对前置已满足节点累积；阻塞节点仅展示原因，不进入强制池。"),
        ("重试导致重复状态变更", "以 run_id、节点 ID 和 row_version 实现 CAS 及幂等键。"),
        ("资源未标注目标节点", "资源规格和发布门校验 target_skill_nodes；缺失时不得计入覆盖。"),
        ("用户失去选择权", "仅自动推荐和排序；创建下一轮任务必须由用户确认。"),
    ], [3000, 6360])

    add_heading(doc, "10. 交付清单", 1)
    for text in [
        "数据库迁移、仓储实现、领域 DTO 与状态机服务。",
        "可解释的多轮选点器及覆盖欠债策略。",
        "生成、发布、反馈链路的状态回写与快照校验。",
        "能力画像与反馈页面的 13 节点进度展示。",
        "单元、集成、迁移、端到端和前端回归测试。",
        "更新后的 API、架构与迁移说明文档。",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "交付完成定义", "所有 13 节点均可被计划、选择、生成、发布、测评与审计；任何节点都能解释其当前状态、是否被前置阻塞、为何被推荐或延后。")

    add_heading(doc, "附录 A. 实施前检查清单", 1)
    add_body(doc, "执行者在开始开发前逐项确认；任何一项未满足时，应先记录风险并与任务负责人对齐，不得以临时绕过方式改变既有公开契约。")
    add_table(doc, ["检查项", "完成条件", "确认"], [
        ("图谱范围", "已确认目标知识库与全部节点 ID；RAG 样例为 13 个能力节点。", "□"),
        ("现状基线", "已确认现有掌握度、资源发布、反馈与生成任务的实际数据模型。", "□"),
        ("迁移方案", "迁移为新增、可回滚、可重复运行；已定义旧资源回填策略。", "□"),
        ("幂等边界", "已定义任务创建、资源发布、反馈提交、重试与取消的幂等键。", "□"),
        ("状态机", "已定义全部合法转移、失败回退和并发 CAS 冲突处理。", "□"),
        ("用户确认", "已确认下一轮只推荐、不自动创建任务的交互约束。", "□"),
        ("测试夹具", "已准备 13 节点、含前置关系、弱项与多轮覆盖的固定 fixture。", "□"),
        ("文档同步", "已确定 API、架构、迁移说明与验收记录的更新责任人。", "□"),
    ], [2600, 5760, 1000])
    add_callout(doc, "启动门槛", "T1 开始前必须完成“图谱范围、现状基线、迁移方案、幂等边界”四项确认；T6 开始前必须完成 T1-T5 的自动化回归。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
